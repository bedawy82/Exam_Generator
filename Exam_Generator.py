"""
Streamlit app: Exam Generator
Features:
- Upload Excel (.xlsx/.xls) or CSV containing questions and options (or a Word .docx with a table)
- Detect correct answer from a 'Correct' column or by reading green cell fill (Excel)
- Upload a bubble-sheet file (xlsx/csv/txt) or paste desired answers for 30 positions
- Generate 6 different Word (.docx) exam models, each with 30 questions
- For each model, reorder choices so the correct answer matches the bubble-sheet mapping
- Download all 6 Word files as a single ZIP with one button

Dependencies:
- streamlit
- pandas
- openpyxl
- python-docx

Run:
pip install streamlit pandas openpyxl python-docx
streamlit run streamlit_exam_generator_app.py

"""

import streamlit as st
import pandas as pd
import random
import io
import zipfile
from docx import Document
from docx.shared import Pt
from typing import List, Tuple, Dict

# For reading excel fill color
from openpyxl import load_workbook

LETTER_ORDER = ["A", "B", "C", "D", "E"]

st.set_page_config(page_title="Exam Models Generator", layout="wide")
st.title(" مولد نماذج الامتحان — Streamlit ")

st.markdown(
    """
    **الوصف:** ارفع ملف الأسئلة (Excel/CSV أو Word table). ثم ارفع نموذج الإجابة (Bubble Sheet) أو ألصق الحروف المطلوبة.
    التطبيق سيختار 30 سؤالًا لكل نموذج ويولّد 6 ملفات Word مع إعادة ترتيب الاختيارات بحيث تطابق نموذج التصحيح.
    """
)

# ----------------------------- Helpers -----------------------------

def read_questions_from_excel_bytes(b: bytes) -> pd.DataFrame:
    # Try pandas first to get dataframe (values only)
    try:
        df = pd.read_excel(io.BytesIO(b), engine="openpyxl")
    except Exception:
        df = pd.read_csv(io.BytesIO(b))
    # Normalize column names
    df.columns = [str(c).strip() for c in df.columns]
    return df


def read_questions_from_docx_bytes(b: bytes) -> pd.DataFrame:
    from docx import Document
    bio = io.BytesIO(b)
    doc = Document(bio)
    # find first table with at least 2 columns
    for tbl in doc.tables:
        if len(tbl.rows) >= 1 and len(tbl.columns) >= 2:
            # read header if present
            data = []
            for r in tbl.rows:
                row = [c.text.strip() for c in r.cells]
                data.append(row)
            # try create df
            try:
                df = pd.DataFrame(data[1:], columns=data[0])
            except Exception:
                df = pd.DataFrame(data)
            return df
    st.error("لم يُعثر على جدول بيانات صالح داخل ملف الوورد.")
    return pd.DataFrame()


def detect_correct_from_excel_bytes(b: bytes, df: pd.DataFrame) -> pd.DataFrame:
    # Use openpyxl to inspect cell fill colors for option columns
    wb = load_workbook(filename=io.BytesIO(b), data_only=True)
    ws = wb.active
    # map header names to column letters
    headers = [str(c).strip() for c in next(ws.iter_rows(min_row=1, max_row=1))]
    # find option columns by header name patterns
    option_cols = []
    for idx, cell in enumerate(next(ws.iter_rows(min_row=1, max_row=1))):
        h = str(cell.value).strip() if cell.value is not None else ""
        if h.lower().startswith("opt") or h.upper() in LETTER_ORDER or any(p in h.lower() for p in ["choice","option","a)","b)"]):
            option_cols.append(idx+1)
    # fallback: consider columns after the question column
    if not option_cols:
        # assume question is first column, options next 4 columns
        option_cols = list(range(2, min(6, ws.max_column+1)))

    # Prepare correct column if not exist
    if 'Correct' not in df.columns and 'correct' not in [c.lower() for c in df.columns]:
        corrects = []
        for row_idx in range(2, 2+len(df)):
            found = None
            for col_idx in option_cols:
                cell = ws.cell(row=row_idx, column=col_idx)
                fill = cell.fill
                try:
                    rgb = fill.start_color.rgb
                except Exception:
                    rgb = None
                if rgb:
                    # Many green fills have 'FF00FF00' or similar. We check presence of '00FF' or '00FF00' or large green component
                    if '00ff00' in str(rgb).lower() or str(rgb).lower().endswith('00ff00') or 'ff00ff00' in str(rgb).lower():
                        # map col idx to letter
                        # determine letter label for this option column by header if available
                        header_cell = ws.cell(row=1, column=col_idx).value
                        label = None
                        if header_cell:
                            header_cell = str(header_cell).strip()
                            # try extract A/B/C
                            for L in LETTER_ORDER:
                                if header_cell.upper().startswith(L):
                                    label = L
                                    break
                        if label is None:
                            # fallback: determine relative index
                            rel = option_cols.index(col_idx)
                            label = LETTER_ORDER[rel] if rel < len(LETTER_ORDER) else LETTER_ORDER[0]
                        found = label
                        break
            corrects.append(found if found else "")
        df['Correct'] = corrects
    return df


def normalize_questions_df(df: pd.DataFrame) -> pd.DataFrame:
    # Ensure columns: Question, OptionA..OptionE, Correct
    cols = list(df.columns)
    # find question column
    qcol = None
    for c in cols:
        if 'question' in c.lower() or 'سؤال' in c.lower() or 'question'==c.lower():
            qcol = c
            break
    if qcol is None:
        qcol = cols[0]
    # find option columns
    option_cols = [c for c in cols if any(p in c.lower() for p in ['option','choice','اختيار','a)','b)','a.','b.']) or c.strip().upper() in LETTER_ORDER]
    if not option_cols:
        # assume next 4 columns after question
        qidx = cols.index(qcol)
        option_cols = cols[qidx+1:qidx+5]
    # rename option columns to OptionA.. based on header or position
    opt_map = {}
    for i, c in enumerate(option_cols):
        opt_map[c] = f'Option{LETTER_ORDER[i]}'
    df = df.rename(columns=opt_map)
    # fill missing Option columns with empty strings
    for i in range(5):
        key = f'Option{LETTER_ORDER[i]}'
        if key not in df.columns:
            df[key] = ""
    # ensure Correct column
    if 'Correct' not in df.columns:
        # try lowercase
        for c in cols:
            if c.lower() == 'correct' or 'الإجابة' in c:
                df = df.rename(columns={c: 'Correct'})
                break
    # standardize Correct to single uppercase letter
    if 'Correct' in df.columns:
        df['Correct'] = df['Correct'].astype(str).str.strip().str.upper().replace({'NAN': ''})
        df['Correct'] = df['Correct'].apply(lambda x: x[0] if x and x[0] in LETTER_ORDER else x)
    return df[[qcol] + [f'Option{L}' for L in LETTER_ORDER if f'Option{L}' in df.columns] + (['Correct'] if 'Correct' in df.columns else [])]


def parse_bubble_sheet(file_bytes: bytes, filename: str) -> List[str]:
    # Accept txt with letters, csv/xlsx with single row/column or a simple list
    name = filename.lower()
    txt = io.BytesIO(file_bytes).getvalue().decode('utf-8', errors='ignore')
    # try CSV
    try:
        if name.endswith('.csv'):
            df = pd.read_csv(io.BytesIO(file_bytes), header=None)
            vals = df.values.flatten()
            letters = [str(v).strip().upper()[0] for v in vals if str(v).strip()]
            return letters
        if name.endswith('.xlsx') or name.endswith('.xls'):
            df = pd.read_excel(io.BytesIO(file_bytes), header=None, engine='openpyxl')
            vals = df.values.flatten()
            letters = [str(v).strip().upper()[0] for v in vals if str(v).strip()]
            return letters
    except Exception:
        pass
    # fallback: take letters from text (A B C ...)
    letters = [tok.strip().upper() for tok in txt.replace(',', ' ').replace('\n', ' ').split() if tok.strip()]
    letters = [t[0] for t in letters if t and t[0] in LETTER_ORDER]
    return letters


def reorder_options_to_match(desired_letter: str, options: List[str], correct_letter: str) -> Tuple[List[str], str]:
    """
    Reorder options list so that the correct option ends up at position corresponding to desired_letter.
    options: list like [optA,optB,optC,optD,...] where their labels correspond to LETTER_ORDER
    correct_letter: the letter (A/B/...) that currently is correct
    returns new_options, new_correct_letter (should equal desired_letter)
    """
    # build mapping letter->index
    letter_to_idx = {L: i for i, L in enumerate(LETTER_ORDER[:len(options)])}
    cur_idx = letter_to_idx.get(correct_letter, 0)
    desired_idx = letter_to_idx.get(desired_letter, 0)
    new_options = options.copy()
    # swap current correct into desired position
    if cur_idx != desired_idx:
        new_options[cur_idx], new_options[desired_idx] = new_options[desired_idx], new_options[cur_idx]
    return new_options, desired_letter


def create_docx_from_questions(questions: List[Dict], title: str = "Exam") -> bytes:
    doc = Document()
    h = doc.add_heading(title, level=1)
    h.alignment = 0
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    for i, q in enumerate(questions, start=1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {q['Question']}").bold = False
        # options
        for idx, opt_text in enumerate(q['Options']):
            doc.add_paragraph(f"{LETTER_ORDER[idx]}. {opt_text}", style='List Number')
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

# ----------------------------- UI -----------------------------

st.header("1) ارفع ملف الأسئلة (Excel/CSV أو Word)")
qfile = st.file_uploader("اختر ملف الأسئلة (.xlsx .csv .docx)", type=["xlsx", "xls", "csv", "docx"] )

st.header("2) ارفع نموذج تصحيح (Bubble Sheet) — ملف بسيط يحتوي على حروف A/B/C لكل سؤال أو الصقها")
bsfile = st.file_uploader("اختر ملف نموذج الإجابة أو اتركه فارغًا لتحديد يدويًا", type=["xlsx","xls","csv","txt","docx"], key='bs')
manual_bs = st.text_area("أو الصق هنا حروف نموذج الإجابة مفصولة بمسافة أو سطر جديد (مثال: A B C ...)")

if qfile is None:
    st.info("ارفع ملف الأسئلة لأجل المتابعة.")
    st.stop()

# Read questions
try:
    if qfile.type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document','application/msword'] or qfile.name.lower().endswith('.docx'):
        df_raw = read_questions_from_docx_bytes(qfile.getvalue())
    else:
        df_raw = read_questions_from_excel_bytes(qfile.getvalue())
except Exception as e:
    st.error(f"فشل قراءة ملف الأسئلة: {e}")
    st.stop()

# If excel, try detect correct from green fill if no Correct column
if qfile.name.lower().endswith(('.xlsx', '.xls')):
    try:
        df_raw = detect_correct_from_excel_bytes(qfile.getvalue(), df_raw)
    except Exception:
        pass

# Normalize
df = normalize_questions_df(df_raw)
st.write("معاينة أول 5 أسطر من ملف الأسئلة بعد التوحيد:")
st.dataframe(df.head())

# Parse bubble sheet
bubble_letters = []
if bsfile is not None:
    try:
        bubble_letters = parse_bubble_sheet(bsfile.getvalue(), bsfile.name)
    except Exception as e:
        st.warning(f"لم أتمكن من قراءة نموذج الإجابة: {e}")
if manual_bs and not bubble_letters:
    bubble_letters = [t[0].upper() for t in manual_bs.split() if t and t[0].upper() in LETTER_ORDER]

if bubble_letters:
    st.write("نموذج الإجابة الذي تم قراءته:")
    st.write(bubble_letters[:30])
else:
    st.info("لم يتم إدخال نموذج الإجابة بعد. يمكنك لصق 30 حرفًا (A/B/C...) في مربع النص أعلاه.")

# Validate we have enough questions
pool_n = len(df)
st.write(f"عدد الأسئلة في الملف: {pool_n}")
if pool_n < 30:
    st.error("الملف يجب أن يحتوي على 30 سؤالاً على الأقل.")
    st.stop()

# Decide sampling strategy
make_disjoint = (pool_n >= 6*30)
st.write("ستُنشأ 6 نماذج مختلفة (كل نموذج 30 سؤال).\n" + ("يمكن إنشاء مجموعات بدون تكرار عبر النماذج لأن عدد الأسئلة كافٍ." if make_disjoint else "سيكون هناك بعض التكرار بين النماذج لأن عدد الأسئلة غير كافٍ لعمل مجموعات منفصلة تمامًا."))

if st.button("إنشاء النماذج الآن"):
    # Build question pool as list of dicts
    # Determine question column name
    qcol = df.columns[0]
    option_cols = [c for c in df.columns if c.startswith('Option')]
    # Ensure Correct exists; if not, try to infer from options where option text contains '(✓)' or '*' etc.
    if 'Correct' not in df.columns:
        # Try infer from option text containing '*' or '(correct)'
        def infer_correct(row):
            for i, c in enumerate(option_cols):
                v = str(row[c])
                if '*' in v or '(correct)' in v.lower() or '✓' in v:
                    return LETTER_ORDER[i]
            return ''
        df['Correct'] = df.apply(infer_correct, axis=1)

    pool = []
    for _, r in df.iterrows():
        opts = [str(r.get(f'Option{L}', '') or '') for L in LETTER_ORDER]
        # trim number of options to those non-empty
        opts = [o for o in opts if o != '']
        correct = str(r.get('Correct', '')).strip().upper()
        # if correct empty, try find which option text equal to some marked text
        pool.append({'Question': str(r[qcol]), 'Options': opts, 'Correct': correct})

    # Generate 6 models
    docs = []
    used_indices = set()
    indices = list(range(len(pool)))
    random.shuffle(indices)
    for model_i in range(6):
        if make_disjoint:
            chosen_idx = indices[model_i*30:(model_i+1)*30]
            if len(chosen_idx) < 30:
                chosen_idx = random.sample(indices, 30)
        else:
            # sample but try to be different
            chosen_idx = random.sample(indices, 30)
        questions_for_doc = []
        for qpos, idx in enumerate(chosen_idx, start=1):
            item = pool[idx]
            options = item['Options']
            correct = item['Correct'] if item['Correct'] else LETTER_ORDER[0]
            # desired letter for this position from bubble sheet if provided, otherwise keep original
            desired_letter = bubble_letters[qpos-1] if qpos-1 < len(bubble_letters) else correct
            # If desired_letter empty, keep current correct
            if not desired_letter:
                desired_letter = correct
            # ensure options length matches LETTER_ORDER subset
            opts_len = len(options)
            letters_for_opts = LETTER_ORDER[:opts_len]
            # If correct not in letters, try locate correct by matching text equals one of options
            if correct not in letters_for_opts:
                # maybe Correct column contains the answer text; try to find index
                corr_text = item.get('Correct')
                found = False
                for ii, opt in enumerate(options):
                    if str(opt).strip() == str(corr_text).strip():
                        correct = letters_for_opts[ii]
                        found = True
                        break
                if not found:
                    correct = letters_for_opts[0]
            # reorder
            new_opts, new_correct = reorder_options_to_match(desired_letter, options, correct)
            questions_for_doc.append({'Question': item['Question'], 'Options': new_opts, 'Correct': new_correct})
        # create docx bytes
        doc_bytes = create_docx_from_questions(questions_for_doc, title=f'Model {model_i+1}')
        docs.append((f'model_{model_i+1}.docx', doc_bytes))

    # zip them
    zip_bio = io.BytesIO()
    with zipfile.ZipFile(zip_bio, mode='w') as zf:
        for name, data in docs:
            zf.writestr(name, data)
    zip_bio.seek(0)

    st.success("تم إنشاء النماذج بنجاح.")
    st.download_button("تحميل النماذج الستة (ZIP)", data=zip_bio.read(), file_name='exam_models.zip', mime='application/zip')

st.markdown("---")
st.write("ملاحظات:")
st.write("- الكود يحاول التعرف على الإجابة الصحيحة من عمود 'Correct' أو من خلية ملوّنة باللون الأخضر في ملف الإكسل.")
st.write("- إذا واجهتك مشكلة في التنسيق، أعد ترتيب الأعمدة أو أضف عمودًا باسم 'Correct' يحتوي على الحرف الصحيح لكل سؤال (A/B/C...).")


